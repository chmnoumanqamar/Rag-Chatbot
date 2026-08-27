"""Automated test suite for Simple RAG Chatbot.

Tests document loaders (TXT, PDF, DOCX, MD), text chunking, hashing,
vector database operations, retrieval, chatbot prompting, and pipeline flow.
"""

import os
import shutil
import tempfile
import unittest
from pathlib import Path

import docx
from pypdf import PdfWriter

from chatbot import Chatbot
from config import CHUNK_OVERLAP, CHUNK_SIZE
from document_loader import Document, DocumentLoader
from rag_pipeline import RAGPipeline
from text_splitter import TextSplitter
from utils import calculate_file_hash, clean_text, validate_file_path
from vector_store import EmbeddingManager, FallbackDeterministicEmbeddings, VectorStoreManager


class TestRAGUtils(unittest.TestCase):
    """Test utility functions for text cleaning, hashing, and validation."""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_clean_text(self):
        raw = "Hello   world!\r\n\r\n\n\nThis is a   test.\xa0With spaces.\n\n\n"
        cleaned = clean_text(raw)
        self.assertNotIn("\r", cleaned)
        self.assertNotIn("\xa0", cleaned)
        self.assertNotIn("   ", cleaned)
        self.assertIn("Hello world!", cleaned)
        self.assertIn("This is a test.", cleaned)

    def test_calculate_file_hash(self):
        test_file = Path(self.temp_dir) / "test.txt"
        test_file.write_text("Unique RAG Content", encoding="utf-8")
        hash1 = calculate_file_hash(test_file)
        self.assertEqual(len(hash1), 64)  # SHA-256 length

        # Same content gives same hash
        hash2 = calculate_file_hash(test_file)
        self.assertEqual(hash1, hash2)

        # Modified content gives different hash
        test_file.write_text("Modified Content", encoding="utf-8")
        hash3 = calculate_file_hash(test_file)
        self.assertNotEqual(hash1, hash3)

    def test_validate_file_path(self):
        test_file = Path(self.temp_dir) / "sample.pdf"
        test_file.write_text("dummy", encoding="utf-8")

        valid, msg = validate_file_path(test_file)
        self.assertTrue(valid)

        # Non-existent
        valid_bad, msg_bad = validate_file_path(Path(self.temp_dir) / "missing.pdf")
        self.assertFalse(valid_bad)

        # Unsupported extension
        unsupported = Path(self.temp_dir) / "data.exe"
        unsupported.write_text("binary", encoding="utf-8")
        valid_un, msg_un = validate_file_path(unsupported)
        self.assertFalse(valid_un)
        self.assertIn("Unsupported file type", msg_un)


class TestDocumentLoader(unittest.TestCase):
    """Test loading and extraction of TXT, MD, DOCX, and PDF documents."""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_load_txt_file(self):
        file_path = Path(self.temp_dir) / "notes.txt"
        file_path.write_text("Python is a powerful language created by Guido van Rossum.", encoding="utf-8")

        docs = DocumentLoader.load_file(file_path)
        self.assertEqual(len(docs), 1)
        self.assertIn("Guido van Rossum", docs[0].content)
        self.assertEqual(docs[0].metadata["file_name"], "notes.txt")
        self.assertEqual(docs[0].metadata["file_type"], ".txt")

    def test_load_pdf_file(self):
        from reportlab.pdfgen import canvas

        file_path = Path(self.temp_dir) / "test_doc.pdf"
        c = canvas.Canvas(str(file_path))
        c.drawString(100, 750, "Natural Language Processing utilizes Transformer models.")
        c.showPage()
        c.save()

        docs = DocumentLoader.load_file(file_path)
        self.assertEqual(len(docs), 1)
        self.assertIn("Transformer models", docs[0].content)
        self.assertEqual(docs[0].metadata["file_type"], ".pdf")
        self.assertEqual(docs[0].metadata["page"], 1)

    def test_load_docx_file(self):
        file_path = Path(self.temp_dir) / "document.docx"
        doc = docx.Document()
        doc.add_heading("Machine Learning Overview", 0)
        doc.add_paragraph("Machine learning is a subset of artificial intelligence.")
        doc.save(str(file_path))

        docs = DocumentLoader.load_file(file_path)
        self.assertEqual(len(docs), 1)
        self.assertIn("subset of artificial intelligence", docs[0].content)
        self.assertEqual(docs[0].metadata["file_type"], ".docx")

    def test_load_empty_file_raises(self):
        file_path = Path(self.temp_dir) / "empty.txt"
        file_path.write_text("   \n\n  ", encoding="utf-8")

        with self.assertRaises(ValueError):
            DocumentLoader.load_file(file_path)


class TestTextSplitter(unittest.TestCase):
    """Test recursive chunking, chunk IDs, and overlap mechanics."""

    def test_short_text_splitting(self):
        splitter = TextSplitter(chunk_size=100, chunk_overlap=20)
        text = "This is a short single sentence."
        chunks = splitter.split_text(text)
        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0], text)

    def test_long_text_splitting_with_overlap(self):
        splitter = TextSplitter(chunk_size=50, chunk_overlap=15)
        text = (
            "Paragraph one is discussing artificial intelligence and its evolution.\n\n"
            "Paragraph two delves into machine learning, deep learning, and neural networks.\n\n"
            "Paragraph three covers retrieval augmented generation architectures."
        )
        chunks = splitter.split_text(text)
        self.assertGreater(len(chunks), 1)
        for c in chunks:
            self.assertLessEqual(len(c), 70)  # Allow slight wiggle for word boundaries

    def test_split_documents_metadata(self):
        splitter = TextSplitter(chunk_size=50, chunk_overlap=10)
        doc = Document(
            content="Sentence one is here. Sentence two is here. Sentence three is here. Sentence four is here.",
            metadata={"source": "test.txt", "file_name": "test.txt", "page": 1},
        )
        chunks = splitter.split_documents([doc])
        self.assertGreater(len(chunks), 1)
        self.assertEqual(chunks[0].metadata["file_name"], "test.txt")
        self.assertEqual(chunks[0].metadata["chunk_id"], 1)
        self.assertEqual(chunks[1].metadata["chunk_id"], 2)


class TestVectorStoreAndRetrieval(unittest.TestCase):
    """Test vector storage, duplicate detection, and similarity search."""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()
        self.embedding_manager = EmbeddingManager(api_key="")  # Uses deterministic embeddings
        self.vector_store = VectorStoreManager(
            db_directory=self.temp_dir,
            embedding_manager=self.embedding_manager,
        )

    def tearDown(self):
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_add_and_search_chunks(self):
        splitter = TextSplitter(chunk_size=200, chunk_overlap=20)
        doc1 = Document(
            content="Python was created by Guido van Rossum and released in 1991.",
            metadata={"source": "python.txt", "file_name": "python.txt", "doc_hash": "hash1"},
        )
        doc2 = Document(
            content="Convolutional Neural Networks are primarily used for computer vision and image processing.",
            metadata={"source": "cnn.txt", "file_name": "cnn.txt", "doc_hash": "hash2"},
        )

        chunks = splitter.split_documents([doc1, doc2])
        count = self.vector_store.add_chunks(chunks)
        self.assertEqual(count, 2)

        stats = self.vector_store.get_stats()
        self.assertEqual(stats["total_chunks"], 2)
        self.assertEqual(stats["total_documents"], 2)

        # Search for Python
        results = self.vector_store.search_similar("Who created Python?", top_k=2, similarity_threshold=0.0)
        self.assertGreater(len(results), 0)
        self.assertEqual(results[0]["metadata"]["file_name"], "python.txt")

    def test_duplicate_detection(self):
        status = self.vector_store.check_document_status("doc.txt", "hash_abc")
        self.assertEqual(status, "new")

        # Add chunk
        doc = Document(content="Sample content", metadata={"file_name": "doc.txt", "doc_hash": "hash_abc"})
        splitter = TextSplitter(chunk_size=100, chunk_overlap=10)
        self.vector_store.add_chunks(splitter.split_documents([doc]))

        # Same hash -> unchanged
        status_same = self.vector_store.check_document_status("doc.txt", "hash_abc")
        self.assertEqual(status_same, "unchanged")

        # Different hash -> modified
        status_mod = self.vector_store.check_document_status("doc.txt", "hash_xyz")
        self.assertEqual(status_mod, "modified")

    def test_clear_database(self):
        doc = Document(content="Temporary data", metadata={"file_name": "temp.txt", "doc_hash": "h1"})
        splitter = TextSplitter(chunk_size=100, chunk_overlap=10)
        self.vector_store.add_chunks(splitter.split_documents([doc]))
        self.assertEqual(self.vector_store.get_stats()["total_chunks"], 1)

        self.vector_store.clear_database()
        self.assertEqual(self.vector_store.get_stats()["total_chunks"], 0)


class TestChatbotAndHistory(unittest.TestCase):
    """Test conversation history and response grounding."""

    def test_history_management(self):
        bot = Chatbot(api_key="")
        for i in range(15):
            bot.add_message("user", f"Question {i}")
            bot.add_message("assistant", f"Answer {i}")

        # MAX_HISTORY is 10 turns -> 20 messages
        self.assertLessEqual(len(bot.history), 20)

        bot.clear_history()
        self.assertEqual(len(bot.history), 0)

    def test_empty_retrieved_chunks_response(self):
        bot = Chatbot(api_key="")
        ans, sources = bot.generate_answer("What is quantum computing?", [])
        self.assertIn("could not find enough relevant information", ans)
        self.assertEqual(len(sources), 0)


class TestRAGPipelineEndToEnd(unittest.TestCase):
    """Test full pipeline workflow: staging -> indexing -> querying -> clearing."""

    def setUp(self):
        self.temp_docs_dir = tempfile.mkdtemp()
        self.temp_vdb_dir = tempfile.mkdtemp()
        self.pipeline = RAGPipeline(
            documents_dir=Path(self.temp_docs_dir),
            vector_db_dir=Path(self.temp_vdb_dir),
        )

    def tearDown(self):
        shutil.rmtree(self.temp_docs_dir, ignore_errors=True)
        shutil.rmtree(self.temp_vdb_dir, ignore_errors=True)

    def test_pipeline_workflow(self):
        # 1. Querying empty KB
        empty_res = self.pipeline.ask("What is machine learning?")
        self.assertIn("Knowledge base is empty", empty_res["answer"])

        # 2. Stage document
        doc_file = Path(self.temp_docs_dir) / "ai_fundamentals.txt"
        doc_file.write_text(
            "Reinforcement Learning is an area of machine learning where agents take actions to maximize reward.",
            encoding="utf-8",
        )
        add_res = self.pipeline.add_document(doc_file)
        self.assertTrue(add_res["success"])

        # 3. Build KB
        build_res = self.pipeline.build_knowledge_base()
        self.assertTrue(build_res["success"])
        self.assertEqual(build_res["indexed_count"], 1)
        self.assertGreater(build_res["total_chunks"], 0)

        # 4. Ask Question
        ask_res = self.pipeline.ask("What is reinforcement learning?")
        self.assertIn("answer", ask_res)
        self.assertGreater(len(ask_res["sources"]), 0)

        # 5. Clear KB
        clear_res = self.pipeline.clear_knowledge_base()
        self.assertTrue(clear_res)
        self.assertEqual(self.pipeline.get_knowledge_base_info()["total_chunks"], 0)


if __name__ == "__main__":
    unittest.main()
