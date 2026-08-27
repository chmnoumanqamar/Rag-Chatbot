"""Document loader module for the Simple RAG Chatbot.

Extracts clean text and metadata from supported document formats (.txt, .pdf, .md, .docx).
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from pypdf import PdfReader
from pypdf.errors import PyPdfError
import docx

from utils import calculate_file_hash, clean_text, logger, validate_file_path


@dataclass
class Document:
    """Represents an extracted document or document page with metadata."""

    content: str
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def source(self) -> str:
        return self.metadata.get("source", "unknown")

    @property
    def file_name(self) -> str:
        return self.metadata.get("file_name", "unknown")

    @property
    def page(self) -> int:
        return self.metadata.get("page", 1)


class DocumentLoader:
    """Handles loading and extracting structured text from various document formats."""

    @classmethod
    def load_file(cls, file_path: str | Path) -> list[Document]:
        """Load a single document from file path and return extracted Document objects.

        Args:
            file_path: Path to the target document.

        Returns:
            List of Document instances (one per page for PDFs, or single for TXT/DOCX/MD).

        Raises:
            FileNotFoundError: If file does not exist.
            ValueError: If file type is unsupported or document is empty.
            RuntimeError: If document extraction fails due to corruption.
        """
        path = Path(file_path).expanduser().resolve()
        is_valid, msg = validate_file_path(path)
        if not is_valid:
            logger.warning("Document validation failed for %s: %s", path.name, msg)
            raise ValueError(msg)

        suffix = path.suffix.lower()
        file_hash = calculate_file_hash(path)

        logger.info("Loading document: %s (%s, hash: %s...)", path.name, suffix, file_hash[:8])

        if suffix in (".txt", ".md"):
            return cls._load_text_file(path, file_hash)
        elif suffix == ".pdf":
            return cls._load_pdf_file(path, file_hash)
        elif suffix == ".docx":
            return cls._load_docx_file(path, file_hash)
        else:
            raise ValueError(
                f"Unsupported file type: {suffix}. Please provide a .txt, .pdf, .md, or .docx file."
            )

    @classmethod
    def _load_text_file(cls, path: Path, file_hash: str) -> list[Document]:
        """Extract text from plain text or markdown file."""
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        content = None

        for enc in encodings:
            try:
                with open(path, "r", encoding=enc) as f:
                    content = f.read()
                break
            except (UnicodeDecodeError, OSError):
                continue

        if content is None:
            raise RuntimeError(f"Could not read text file '{path.name}' with standard encodings.")

        cleaned = clean_text(content)
        if not cleaned:
            raise ValueError(f"Document '{path.name}' is empty or contains only whitespace.")

        metadata = {
            "source": str(path),
            "file_name": path.name,
            "file_type": path.suffix.lower(),
            "doc_hash": file_hash,
            "page": 1,
            "char_count": len(cleaned),
        }
        return [Document(content=cleaned, metadata=metadata)]

    @classmethod
    def _load_pdf_file(cls, path: Path, file_hash: str) -> list[Document]:
        """Extract text page-by-page from a PDF file using pypdf."""
        documents: list[Document] = []
        try:
            reader = PdfReader(str(path))
        except PyPdfError as e:
            logger.error("PyPdf error while reading %s: %s", path.name, e)
            raise RuntimeError(f"Failed to parse PDF '{path.name}': Corrupt or invalid format.") from e
        except Exception as e:
            logger.error("Unexpected error reading PDF %s: %s", path.name, e)
            raise RuntimeError(f"Failed to open PDF '{path.name}': {e}") from e

        if reader.is_encrypted:
            try:
                # Attempt decrypt with empty password
                reader.decrypt("")
            except Exception as e:
                raise RuntimeError(
                    f"PDF '{path.name}' is password-protected and cannot be read."
                ) from e

        total_pages = len(reader.pages)
        if total_pages == 0:
            raise ValueError(f"PDF '{path.name}' contains no pages.")

        total_extracted_chars = 0
        for page_idx, page in enumerate(reader.pages, start=1):
            try:
                raw_text = page.extract_text() or ""
            except Exception as e:
                logger.warning("Error extracting text from page %d in %s: %s", page_idx, path.name, e)
                raw_text = ""

            cleaned = clean_text(raw_text)
            if cleaned:
                total_extracted_chars += len(cleaned)
                metadata = {
                    "source": str(path),
                    "file_name": path.name,
                    "file_type": ".pdf",
                    "doc_hash": file_hash,
                    "page": page_idx,
                    "total_pages": total_pages,
                    "char_count": len(cleaned),
                }
                documents.append(Document(content=cleaned, metadata=metadata))

        if not documents or total_extracted_chars == 0:
            raise ValueError(
                f"PDF '{path.name}' does not contain any extractable text. "
                "It may contain scanned images without OCR."
            )

        logger.info(
            "Successfully extracted %d pages (%d characters) from PDF %s",
            len(documents),
            total_extracted_chars,
            path.name,
        )
        return documents

    @classmethod
    def _load_docx_file(cls, path: Path, file_hash: str) -> list[Document]:
        """Extract text from Word document (.docx)."""
        try:
            doc = docx.Document(str(path))
        except Exception as e:
            logger.error("Error opening DOCX %s: %s", path.name, e)
            raise RuntimeError(f"Failed to parse Word document '{path.name}': {e}") from e

        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        raw_content = "\n\n".join(full_text)
        cleaned = clean_text(raw_content)

        if not cleaned:
            raise ValueError(f"Word document '{path.name}' contains no readable text.")

        metadata = {
            "source": str(path),
            "file_name": path.name,
            "file_type": ".docx",
            "doc_hash": file_hash,
            "page": 1,
            "char_count": len(cleaned),
        }
        return [Document(content=cleaned, metadata=metadata)]

    @classmethod
    def load_directory(cls, dir_path: str | Path) -> list[Document]:
        """Load all supported documents in a directory."""
        path = Path(dir_path).expanduser().resolve()
        if not path.is_dir():
            raise NotADirectoryError(f"Directory not found: {dir_path}")

        all_docs: list[Document] = []
        for file in sorted(path.iterdir()):
            if file.is_file() and file.suffix.lower() in (".txt", ".pdf", ".md", ".docx"):
                try:
                    docs = cls.load_file(file)
                    all_docs.extend(docs)
                except Exception as e:
                    logger.warning("Skipping file '%s' due to error: %s", file.name, e)

        return all_docs
